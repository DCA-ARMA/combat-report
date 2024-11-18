# templates.py

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import os
import re
from bidi.algorithm import get_display
from matplotlib import rcParams

def generate_word_document(sections, output_path, date="", signature="", title="", grades_data=None):
    """
    Generates a Word document from the given sections dictionary, including title, date, signature, and grades.
    """
    document = Document()

    # Set default font to support Hebrew characters
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Ensure this font supports Hebrew
    font.size = Pt(12)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')  # Fix for Hebrew characters

    # --- Cover Page ---
    # Add a section with a different page setup
    section = document.sections[-1]
    section.start_type = WD_SECTION.NEW_PAGE

    # Add DCA logo
    logo_path = "dca_logo.png"  # Ensure this file exists in your directory
    if os.path.exists(logo_path):
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(2))
    else:
        print(f"Warning: Logo file '{logo_path}' not found. Skipping logo.")

    # Add the title
    cover_title = document.add_heading("דוח סיכום קרב", level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title_run = cover_title.runs[0]
    cover_title_run.font.name = 'Arial'
    cover_title_run.font.size = Pt(24)
    cover_title_run.font.bold = True
    cover_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add the date
    if date:
        date_paragraph = document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(date)
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(14)
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add a brief description of DCA
    dca_description = (
        "DCA היא חברה המתמחה באימונים בסימולטורים מתקדמים עבור כוחות הביטחון. "
        "אנו מחויבים להעניק את ההכשרה המקצועית והאיכותית ביותר למתאמנים שלנו."
    )
    description_paragraph = document.add_paragraph()
    description_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description_run = description_paragraph.add_run(dca_description)
    description_run.font.name = 'Arial'
    description_run.font.size = Pt(14)
    description_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add a page break after the cover page
    document.add_page_break()

    # --- Existing Content (Sections) ---
    # Add the title for the main content
    if title:
        title_paragraph = document.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
        title_run = title_paragraph.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add each section to the document
    for section_title, hebrew_title in [("Introduction", "מבוא"), ("Scenario 1", "תרחיש 1"), ("Scenario 2", "תרחיש 2"), ("Summary", "סיכום")]:
        if sections[section_title]:
            # Add section title
            heading = document.add_heading(hebrew_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align Right
            heading_run = heading.runs[0]
            heading_run.font.name = 'Arial'
            heading_run.font.size = Pt(16)
            heading_run.font.bold = True
            heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            # Add section content
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align Right
            content = sections[section_title]
            # Remove '#' and '*' from content
            content = content.replace('#', '').replace('*', '')
            run = paragraph.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # --- Grades Section ---
    if grades_data:
        # Add a page break before grades
        document.add_page_break()

        # Add Grades Section
        grades_heading = document.add_heading('דוח ציונים', level=1)
        grades_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        grades_heading_run = grades_heading.runs[0]
        grades_heading_run.font.name = 'Arial'
        grades_heading_run.font.size = Pt(16)
        grades_heading_run.font.bold = True
        grades_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Insert graphs only, without free text grades
        for part_name, part_data in grades_data.items():
            if part_name == "final_grade":
                continue
            # Add part heading
            part_heading = document.add_heading(part_name, level=2)
            part_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            part_heading_run = part_heading.runs[0]
            part_heading_run.font.name = 'Arial'
            part_heading_run.font.size = Pt(14)
            part_heading_run.font.bold = True
            part_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            # Generate and insert graph for the part
            # Sanitize the part name for the filename
            sanitized_part_name = sanitize_filename(part_name)
            chart_path = f"{sanitized_part_name}_chart.png"
            create_bar_chart(part_data['items'], part_name, chart_path)
            document.add_picture(chart_path, width=Inches(6))
            # Remove the chart image file after inserting
            os.remove(chart_path)

        # Add final grade
        final_grade_paragraph = document.add_paragraph()
        final_grade_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = final_grade_paragraph.add_run(f"ציון סופי: {grades_data['final_grade']}")
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Generate and insert final grade graph
        final_chart_path = "final_grade_chart.png"
        create_final_grade_chart(grades_data, final_chart_path)
        document.add_picture(final_chart_path, width=Inches(6))
        # Remove the chart image file after inserting
        os.remove(final_chart_path)

    # --- Final Page ---
    # Add a page break before the final page
    document.add_page_break()

    # Add final page content
    thank_you_title = document.add_heading("תודה שהשתתפתם באימון שלנו", level=0)
    thank_you_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thank_you_title_run = thank_you_title.runs[0]
    thank_you_title_run.font.name = 'Arial'
    thank_you_title_run.font.size = Pt(24)
    thank_you_title_run.font.bold = True
    thank_you_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Thank you message
    thank_you_message = (
        "אנו מודים לכם על השתתפותכם באימון שלנו. "
        "נשמח לעמוד לשירותכם בכל עת."
    )
    message_paragraph = document.add_paragraph()
    message_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    message_run = message_paragraph.add_run(thank_you_message)
    message_run.font.name = 'Arial'
    message_run.font.size = Pt(14)
    message_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Contact information
    contact_info = {
        "אתר האינטרנט": "https://www.dca.co.il",
        "דוא\"ל": "contact@dca.co.il",
        "טלפון": "+972-3-1234567",
        "LinkedIn": "https://www.linkedin.com/company/dca-israel/"
    }

    for label, info in contact_info.items():
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_paragraph.add_run(f"{label}: {info}")
        contact_run.font.name = 'Arial'
        contact_run.font.size = Pt(12)
        contact_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        # Underline the links
        if label != "טלפון":
            contact_run.font.underline = True
        # Add hyperlink (Note: Word doesn't support hyperlinks added via python-docx in a straightforward way)

    # Add signature at the end
    if signature:
        document.add_paragraph()  # Add empty paragraph for spacing
        signature_paragraph = document.add_paragraph()
        signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = signature_paragraph.add_run(signature)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Save the document
    document.save(output_path)

def create_bar_chart(items, title, output_path):
    """
    Creates a bar chart for the given items and saves it to the specified path.
    """
    # Reverse the items for right-to-left display
    labels = list(items.keys())[::-1]
    grades = list(items.values())[::-1]

    # Use get_display to handle RTL text
    labels = [get_display(label) for label in labels]
    title = get_display(title)

    # Set font to Arial (or another font that supports Hebrew)
    rcParams['font.family'] = 'Arial'

    plt.figure(figsize=(8, 4))
    bars = plt.barh(labels, grades, color='skyblue')
    plt.xlabel(get_display("ציון"), fontsize=12)
    plt.title(title, fontsize=14)
    plt.xlim(0, 10)
    plt.tight_layout()
    # Add grade labels next to the bars
    for bar, grade in zip(bars, grades):
        plt.text(grade + 0.1, bar.get_y() + bar.get_height()/2, f'{grade}', va='center')
    plt.savefig(output_path)
    plt.close()

def create_final_grade_chart(grades_data, output_path):
    """
    Creates a bar chart for the final grades of each part.
    """
    parts = [part for part in grades_data if part != 'final_grade']
    averages = [grades_data[part]['average'] for part in parts]

    # Reverse the parts and averages for right-to-left display
    parts = parts[::-1]
    averages = averages[::-1]

    # Use get_display to handle RTL text
    parts = [get_display(part) for part in parts]
    title = get_display('ציון ממוצע לכל חלק')

    # Set font to Arial (or another font that supports Hebrew)
    rcParams['font.family'] = 'Arial'

    plt.figure(figsize=(8, 4))
    bars = plt.barh(parts, averages, color='lightgreen')
    plt.xlabel(get_display("ציון ממוצע"), fontsize=12)
    plt.title(title, fontsize=14)
    plt.xlim(0, 10)
    plt.tight_layout()
    # Add average labels next to the bars
    for bar, avg in zip(bars, averages):
        plt.text(avg + 0.1, bar.get_y() + bar.get_height()/2, f'{avg}', va='center')
    plt.savefig(output_path)
    plt.close()

def sanitize_filename(filename):
    """
    Sanitizes the filename by removing or replacing invalid characters.
    """
    # Define invalid characters for Windows filenames
    invalid_chars = r'[<>:"/\\|?*]'
    # Replace invalid characters with an underscore or remove them
    sanitized = re.sub(invalid_chars, '', filename)
    return sanitized
